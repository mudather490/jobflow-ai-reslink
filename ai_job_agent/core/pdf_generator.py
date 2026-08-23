import os
import re
from pathlib import Path
from typing import Tuple, List, Optional, Dict, Any
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.pdfgen import canvas

from core.resume_parser import UserProfile
from core.template_registry import get_template, ResumeTemplate
from config import OUTPUT_DIR


class NumberedCanvas(canvas.Canvas):
    """
    Two-pass canvas to calculate total page count and draw clean page numbers in bottom-right.
    """
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            super().showPage()
        super().save()

    def draw_page_number(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.HexColor("#475569"))
        self.drawRightString(letter[0] - 40, 28, f"{self._pageNumber}")
        self.restoreState()


def _hex_to_rgb(hex_str: str) -> RGBColor:
    hex_clean = hex_str.lstrip("#")
    r = int(hex_clean[0:2], 16)
    g = int(hex_clean[2:4], 16)
    b = int(hex_clean[4:6], 16)
    return RGBColor(r, g, b)


class ResumeDocumentGenerator:
    """
    Multi-Template Resume Document Engine.
    Generates exact matching DOCX and PDF documents across 4 world-class professional templates:
    1. Modern (Royal Blue)
    2. Harvard Consulting (MBB Deep Navy)
    3. Corporate Elite (Deep Navy & Gold)
    4. Tech Specialist (Purple & Cyan)
    """

    @staticmethod
    def _sanitize_filename(text: str) -> str:
        cleaned = re.sub(r"[^\w\-_\. ]", "", text).replace(" ", "_")
        cleaned = re.sub(r"_+", "_", cleaned).strip("_-.")
        return cleaned or "Resume"

    @classmethod
    def generate_docx(cls, profile: UserProfile, output_path: str, template_id: str = "modern") -> str:
        """
        Creates an executive DOCX resume matching the chosen template style.
        """
        tmpl = get_template(template_id)
        prim_rgb = _hex_to_rgb(tmpl.primary_color)
        sec_rgb = _hex_to_rgb(tmpl.secondary_color)

        doc = docx.Document()

        # Set official typography based on template standard
        if template_id == "harvard_consulting":
            doc.styles['Normal'].font.name = 'Times New Roman'
        elif template_id == "tech_specialist":
            doc.styles['Normal'].font.name = 'Segoe UI'
        elif template_id == "corporate_elite":
            doc.styles['Normal'].font.name = 'Calibri'
        else: # modern (Google FAANG)
            doc.styles['Normal'].font.name = 'Arial'

        for section in doc.sections:
            section.top_margin = Inches(0.55)
            section.bottom_margin = Inches(0.55)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        # Header rendering based on template header style
        if tmpl.header_style == "centered_classic":
            # Harvard Consulting centered header
            p_name = doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_name = p_name.add_run(profile.full_name.upper())
            r_name.bold = True
            r_name.font.size = Pt(17)
            r_name.font.color.rgb = prim_rgb

            p_head = doc.add_paragraph()
            p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            head_text = profile.headline or profile.target_role or ""
            r_head = p_head.add_run(head_text)
            r_head.font.size = Pt(9.5)
            r_head.font.color.rgb = sec_rgb

            contacts = []
            if profile.contact.phone: contacts.append(profile.contact.phone)
            if profile.contact.email: contacts.append(profile.contact.email.replace("mailto:", ""))
            if profile.contact.location: contacts.append(profile.contact.location)
            if profile.contact.linkedin: contacts.append(profile.contact.linkedin)
            if profile.contact.github: contacts.append(profile.contact.github)

            p_cnt = doc.add_paragraph()
            p_cnt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cnt.paragraph_format.space_after = Pt(6)
            r_cnt = p_cnt.add_run(" | ".join(contacts))
            r_cnt.font.size = Pt(8.5)
            r_cnt.font.color.rgb = RGBColor(71, 85, 105)

        else:
            # 2-Column Table Header (Modern, Corporate Elite, Tech Specialist)
            t_header = doc.add_table(rows=1, cols=2)
            t_header.autofit = True
            c_left, c_right = t_header.rows[0].cells

            p_name = c_left.paragraphs[0]
            r_name = p_name.add_run(profile.full_name.upper())
            r_name.bold = True
            r_name.font.size = Pt(18)
            r_name.font.color.rgb = prim_rgb

            p_head = c_left.add_paragraph()
            head_text = profile.headline or profile.target_role or ""
            r_head = p_head.add_run(head_text)
            r_head.font.size = Pt(10)
            r_head.font.color.rgb = prim_rgb

            p_contact = c_right.paragraphs[0]
            p_contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if profile.contact.phone:
                r_ph = p_contact.add_run(f"Phone: {profile.contact.phone}\n")
                r_ph.font.size = Pt(8.5)
            if profile.contact.email:
                r_em = p_contact.add_run(f"Email: {profile.contact.email.replace('mailto:', '')}\n")
                r_em.font.size = Pt(8.5)
            if profile.contact.location:
                r_loc = p_contact.add_run(f"Location: {profile.contact.location}\n")
                r_loc.font.size = Pt(8.5)
            if profile.contact.linkedin:
                r_li = p_contact.add_run(f"LinkedIn: {profile.contact.linkedin}\n")
                r_li.font.size = Pt(8.5)
            if profile.contact.github:
                r_gh = p_contact.add_run(f"GitHub: {profile.contact.github}")
                r_gh.font.size = Pt(8.5)

        def add_heading(title: str, upper: bool = False):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(title.upper() if upper else title)
            r.bold = True
            r.font.size = Pt(11.5 if not upper else 10.5)
            r.font.color.rgb = prim_rgb

        # Helper for skills block
        def render_skills():
            if hasattr(profile, 'categorized_skills') and any(profile.categorized_skills.values()):
                add_heading("TECHNICAL SKILLS", upper=True)
                for cat, sks in profile.categorized_skills.items():
                    if sks:
                        p_sk = doc.add_paragraph(style="List Bullet")
                        r_cat = p_sk.add_run(f"{cat}: ")
                        r_cat.bold = True
                        r_cat.font.color.rgb = prim_rgb
                        p_sk.add_run(", ".join(sks))
                        p_sk.paragraph_format.space_after = Pt(2)
            elif profile.skills:
                add_heading("TECHNICAL SKILLS", upper=True)
                p_sk = doc.add_paragraph(style="List Bullet")
                r_cat = p_sk.add_run("Skills & Core Competencies: ")
                r_cat.bold = True
                r_cat.font.color.rgb = prim_rgb
                p_sk.add_run(", ".join(profile.skills))

        # Helper for projects block
        def render_projects():
            if profile.projects:
                add_heading("Production & Practical Projects", upper=True)
                for proj in profile.projects:
                    p_p = doc.add_paragraph(style="List Bullet")
                    r_pn = p_p.add_run(proj.name.strip(" :—"))
                    r_pn.bold = True
                    if proj.technologies:
                        p_p.add_run(f" | {', '.join(proj.technologies)}")
                    if proj.repository:
                        p_p.add_run(f" | {proj.repository.strip()}")

                    if proj.subtitle:
                        p_sub = doc.add_paragraph()
                        r_sub = p_sub.add_run(proj.subtitle)
                        r_sub.bold = True

                    if proj.bullets:
                        for b in proj.bullets:
                            p_b = doc.add_paragraph(style="List Bullet")
                            p_b.add_run(b)
                    elif proj.description:
                        p_d = doc.add_paragraph()
                        p_d.add_run(proj.description)

        # Helper for experience block
        def render_experience():
            if profile.experience:
                add_heading("Professional Experience", upper=True)
                for exp in profile.experience:
                    p_r = doc.add_paragraph()
                    clean_role = exp.role
                    if clean_role.lower() in ["professional role", "work professional", "lead professional", "professional profile", "role"]:
                        clean_role = profile.headline.split("|")[0].strip() if profile.headline else "Software Engineer"
                    exp_title = clean_role
                    if exp.company and exp.company not in ["Independent", "General", ""]:
                        exp_title = f"{clean_role} — {exp.company}"
                    if exp.duration:
                        exp_title = f"{exp_title} | {exp.duration}"
                    r_ro = p_r.add_run(exp_title)
                    r_ro.bold = True
                    r_ro.font.color.rgb = prim_rgb

                    if exp.subtitle:
                        p_sub = doc.add_paragraph()
                        r_sub = p_sub.add_run(exp.subtitle)
                        r_sub.bold = True
                        r_sub.font.color.rgb = prim_rgb

                    if exp.summary:
                        doc.add_paragraph(exp.summary)

                    if exp.bullets:
                        for b in exp.bullets:
                            p_b = doc.add_paragraph(style="List Bullet")
                            p_b.add_run(b)

        # Helper for education & certs
        def render_education_and_certs():
            if profile.certifications:
                add_heading("Certifications & Continuous Education", upper=True)
                for cert in profile.certifications:
                    p_c = doc.add_paragraph(style="List Bullet")
                    r_cn = p_c.add_run(f"{cert.name}")
                    r_cn.bold = True
                    if cert.issuer:
                        p_c.add_run(f" — {cert.issuer}")
                    if cert.status:
                        p_c.add_run(f": {cert.status}")
                    if cert.details:
                        p_c.add_run(f" — {cert.details}")

            if profile.education:
                add_heading("Education / Transition Background", upper=True)
                for edu in profile.education:
                    p_e = doc.add_paragraph(style="List Bullet")
                    r_en = p_e.add_run(f"{edu.institution}")
                    r_en.bold = True
                    if edu.degree and edu.degree != "Degree":
                        p_e.add_run(f" — {edu.degree}")
                    if edu.year:
                        p_e.add_run(f" | {edu.year}")
                    if edu.details:
                        p_ed = doc.add_paragraph()
                        p_ed.add_run(edu.details)

        # Template Flow Sequencing
        if tmpl.id == "corporate_elite":
            if profile.summary:
                p_sum = doc.add_paragraph(profile.summary)
                p_sum.paragraph_format.space_after = Pt(4)
            render_skills()
            render_projects()
            render_experience()
            render_education_and_certs()

        elif tmpl.id == "tech_specialist":
            if profile.summary:
                p_sum = doc.add_paragraph(profile.summary)
                p_sum.paragraph_format.space_after = Pt(4)
            render_skills()
            render_projects()
            render_experience()
            render_education_and_certs()

        elif tmpl.id == "harvard_consulting":
            if profile.summary:
                p_sum = doc.add_paragraph(profile.summary)
                p_sum.paragraph_format.space_after = Pt(4)
            render_experience()
            render_education_and_certs()
            render_skills()
            render_projects()

        else: # modern default
            if profile.summary:
                p_sum = doc.add_paragraph(profile.summary)
                p_sum.paragraph_format.space_after = Pt(4)
            render_education_and_certs()
            render_skills()
            render_projects()
            render_experience()

        if profile.additional_background:
            add_heading("ADDITIONAL BACKGROUND :", upper=True)
            doc.add_paragraph(profile.additional_background)

        try:
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            return output_path
        except (OSError, PermissionError) as pe:
            import tempfile
            temp_out = str(Path(tempfile.gettempdir()) / Path(output_path).name)
            doc.save(temp_out)
            return temp_out

    @classmethod
    def generate_docx_bytes(cls, profile: UserProfile, template_id: str = "modern") -> Any:
        """
        Generates the professional Word .docx document directly in RAM (io.BytesIO)
        without relying on local disk file reads or writes. 100% Vercel Serverless Immune!
        """
        import io
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        
        # Set 0.75-inch margins (standard ATS formatting)
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        tmpl = get_template(template_id)
        prim_rgb = RGBColor.from_string(tmpl.primary_color.lstrip('#'))

        # Header: Candidate Name
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_name = p_name.add_run(profile.full_name or "MUDATHER MOHAMMED")
        r_name.bold = True
        r_name.font.size = Pt(18)
        r_name.font.color.rgb = prim_rgb

        # Contact Line
        if profile.contact:
            p_contact = doc.add_paragraph()
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c_parts = []
            if profile.contact.location: c_parts.append(profile.contact.location)
            if profile.contact.email: c_parts.append(profile.contact.email)
            if profile.contact.phone: c_parts.append(profile.contact.phone)
            if profile.contact.linkedin: c_parts.append(profile.contact.linkedin)
            if profile.contact.github: c_parts.append(profile.contact.github)
            r_contact = p_contact.add_run(" • ".join(c_parts))
            r_contact.font.size = Pt(9.5)
            r_contact.font.color.rgb = RGBColor(100, 100, 100)

        # Helper for Headings
        def add_sec_heading(title: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(title.upper())
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = prim_rgb

        # Summary
        if profile.summary:
            p_sum = doc.add_paragraph(profile.summary)
            p_sum.paragraph_format.space_after = Pt(4)

        # Technical Skills
        if profile.skills:
            add_sec_heading("Technical Skills & Core Competencies")
            p_sk = doc.add_paragraph(style="List Bullet")
            r_sk = p_sk.add_run("Skills: ")
            r_sk.bold = True
            r_sk.font.color.rgb = prim_rgb
            p_sk.add_run(", ".join(profile.skills))

        # Professional Experience
        if profile.experience:
            add_sec_heading("Professional Experience")
            for exp in profile.experience:
                p_exp = doc.add_paragraph()
                r_r = p_exp.add_run(f"{exp.role} — {exp.company}")
                r_r.bold = True
                if exp.duration:
                    r_d = p_exp.add_run(f"\t{exp.duration}")
                    r_d.italic = True
                    p_exp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))

                if exp.bullets:
                    for b in exp.bullets:
                        doc.add_paragraph(f"• {b}", style="List Bullet")

        # Projects
        if profile.projects:
            add_sec_heading("Production Projects")
            for proj in profile.projects:
                p_p = doc.add_paragraph()
                r_p = p_p.add_run(proj.name)
                r_p.bold = True
                if proj.technologies:
                    p_p.add_run(f" ({', '.join(proj.technologies)})")
                if proj.bullets:
                    for b in proj.bullets:
                        doc.add_paragraph(f"• {b}", style="List Bullet")

        # Save directly to RAM BytesIO stream
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream



    @classmethod
    def generate_pdf(cls, profile: UserProfile, output_path: str, template_id: str = "modern") -> str:
        """
        Creates a high-fidelity PDF adhering strictly to the chosen professional template.
        """
        tmpl = get_template(template_id)
        prim_col = colors.HexColor(tmpl.primary_color)
        sec_col = colors.HexColor(tmpl.secondary_color)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=36,
            leftMargin=36,
            topMargin=32,
            bottomMargin=36,
        )

        font_normal = tmpl.font_name if tmpl.font_name in ["Helvetica", "Times-Roman", "Courier"] else "Helvetica"
        font_bold = tmpl.font_bold if tmpl.font_bold in ["Helvetica-Bold", "Times-Bold", "Courier-Bold"] else "Helvetica-Bold"
        font_italic = tmpl.font_oblique if tmpl.font_oblique in ["Helvetica-Oblique", "Times-Italic", "Courier-Oblique"] else "Helvetica-Oblique"

        styles = getSampleStyleSheet()

        name_style = ParagraphStyle(
            "UserName",
            parent=styles["Normal"],
            fontName=font_bold,
            fontSize=20 if tmpl.header_style != "centered_classic" else 18,
            leading=24,
            textColor=prim_col,
            alignment=1 if tmpl.header_style == "centered_classic" else 0,
        )

        headline_style = ParagraphStyle(
            "UserHeadline",
            parent=styles["Normal"],
            fontName=font_normal,
            fontSize=10,
            leading=13.5,
            textColor=prim_col if tmpl.id not in ["corporate_elite", "minimal"] else sec_col,
            alignment=1 if tmpl.header_style == "centered_classic" else 0,
        )

        contact_right_style = ParagraphStyle(
            "ContactRight",
            parent=styles["Normal"],
            fontName=font_normal,
            fontSize=8.5,
            leading=12,
            textColor=colors.HexColor("#1E293B"),
            alignment=2,
        )

        contact_center_style = ParagraphStyle(
            "ContactCenter",
            parent=styles["Normal"],
            fontName=font_normal,
            fontSize=8.5,
            leading=12,
            textColor=colors.HexColor("#334155"),
            alignment=1,
        )

        summary_style = ParagraphStyle(
            "SummaryText",
            parent=styles["Normal"],
            fontName=font_normal,
            fontSize=9.0,
            leading=13.0,
            textColor=colors.HexColor("#1E293B"),
        )

        section_title_style = ParagraphStyle(
            "SectionTitle",
            parent=styles["Normal"],
            fontName=font_bold,
            fontSize=12.0 if tmpl.id != "harvard" else 11.5,
            leading=15,
            textColor=prim_col,
            spaceBefore=6,
            spaceAfter=2.5,
        )

        bullet_style = ParagraphStyle(
            "BulletText",
            parent=styles["Normal"],
            fontName=font_normal,
            fontSize=8.7,
            leading=12.2,
            textColor=colors.HexColor("#1E293B"),
            leftIndent=12,
            firstLineIndent=-8,
            spaceAfter=1.8,
        )

        sub_bullet_style = ParagraphStyle(
            "SubBulletText",
            parent=styles["Normal"],
            fontName=font_normal,
            fontSize=8.7,
            leading=12.2,
            textColor=colors.HexColor("#1E293B"),
            leftIndent=12,
            spaceAfter=1.5,
        )

        role_title_style = ParagraphStyle(
            "RoleTitle",
            parent=styles["Normal"],
            fontName=font_bold,
            fontSize=9.5,
            leading=13.5,
            textColor=prim_col,
            spaceBefore=3.5,
            spaceAfter=1.5,
        )

        elements = []

        # 1. Header Logic
        if tmpl.header_style == "centered_classic":
            # Centered Harvard MBB Header
            elements.append(Paragraph(f"<b>{profile.full_name.upper()}</b>", name_style))
            elements.append(Spacer(1, 2))
            head_text = profile.headline or profile.target_role or ""
            if head_text:
                elements.append(Paragraph(head_text, headline_style))
                elements.append(Spacer(1, 3))

            cnt_items = []
            if profile.contact.phone: cnt_items.append(f"<b>Phone:</b> {profile.contact.phone}")
            if profile.contact.email:
                em = profile.contact.email.replace("mailto:", "")
                cnt_items.append(f'<a href="mailto:{em}" color="{tmpl.primary_color}"><u>{em}</u></a>')
            if profile.contact.location: cnt_items.append(profile.contact.location)
            if profile.contact.linkedin:
                li_url = profile.contact.linkedin if profile.contact.linkedin.startswith("http") else f"https://{profile.contact.linkedin}"
                cnt_items.append(f'<a href="{li_url}" color="{tmpl.primary_color}"><u>LinkedIn</u></a>')
            if profile.contact.github:
                gh_url = profile.contact.github if profile.contact.github.startswith("http") else f"https://{profile.contact.github}"
                cnt_items.append(f'<a href="{gh_url}" color="{tmpl.primary_color}"><u>GitHub</u></a>')

            elements.append(Paragraph(" &nbsp;•&nbsp; ".join(cnt_items), contact_center_style))
            elements.append(Spacer(1, 4))
            elements.append(HRFlowable(width="100%", thickness=1, color=prim_col, spaceBefore=2, spaceAfter=6))

        else:
            # 2-Column Header
            left_header = [
                Paragraph(f"<b>{profile.full_name.upper()}</b>", name_style),
                Spacer(1, 2),
            ]
            head_text = profile.headline or profile.target_role or ""
            if head_text:
                left_header.append(Paragraph(head_text, headline_style))

            right_header = []
            if profile.contact.phone:
                right_header.append(Paragraph(f'<b>Phone:</b> {profile.contact.phone}', contact_right_style))
            if profile.contact.email:
                em = profile.contact.email.replace("mailto:", "")
                right_header.append(Paragraph(f'<b>Email:</b> <a href="mailto:{em}" color="{tmpl.primary_color}"><u>{em}</u></a>', contact_right_style))
            if profile.contact.location:
                right_header.append(Paragraph(f'<b>Location:</b> {profile.contact.location}', contact_right_style))
            if profile.contact.linkedin:
                li_url = profile.contact.linkedin if profile.contact.linkedin.startswith("http") else f"https://{profile.contact.linkedin}"
                right_header.append(Paragraph(f'<b>LinkedIn:</b> <a href="{li_url}" color="{tmpl.primary_color}"><u>{li_url}</u></a>', contact_right_style))
            if profile.contact.github:
                gh_url = profile.contact.github if profile.contact.github.startswith("http") else f"https://{profile.contact.github}"
                right_header.append(Paragraph(f'<b>GitHub:</b> <a href="{gh_url}" color="{tmpl.primary_color}"><u>{gh_url}</u></a>', contact_right_style))

            header_table = Table([[left_header, right_header]], colWidths=[290, 250])
            header_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ]))
            elements.append(header_table)

            if tmpl.id == "corporate_elite":
                elements.append(Spacer(1, 4))
                elements.append(HRFlowable(width="100%", thickness=1.5, color=sec_col, spaceBefore=2, spaceAfter=6))
            else:
                elements.append(Spacer(1, 6))

        # Reusable Section Header Builder with Signature Template Divider Bar
        def add_pdf_heading(title: str, upper: bool = False):
            t_text = title.upper() if (upper or tmpl.id in ["harvard_consulting", "tech_specialist", "corporate_elite"]) else title
            elements.append(Spacer(1, 4))
            elements.append(Paragraph(f"<b>{t_text}</b>", section_title_style))
            
            # Signature divider bars per template
            if tmpl.id == "harvard_consulting":
                elements.append(HRFlowable(width="100%", thickness=1.5, color=prim_col, spaceBefore=1, spaceAfter=4))
            elif tmpl.id == "corporate_elite":
                elements.append(HRFlowable(width="100%", thickness=1.2, color=sec_col, spaceBefore=1, spaceAfter=4))
            elif tmpl.id == "tech_specialist":
                elements.append(HRFlowable(width="100%", thickness=1.2, color=prim_col, spaceBefore=1, spaceAfter=4))
            else: # modern
                elements.append(HRFlowable(width="100%", thickness=0.8, color=prim_col, spaceBefore=1, spaceAfter=4))

        # Reusable Section Renderers
        def add_pdf_education():
            if profile.education:
                add_pdf_heading("Education / Transition Background", upper=(tmpl.id != "modern"))
                for edu in profile.education:
                    yr_str = f" | {edu.year}" if edu.year else ""
                    deg_str = f" — {edu.degree}" if edu.degree and edu.degree != "Degree" else ""
                    elements.append(Paragraph(f"• <b>{edu.institution}{deg_str}</b>{yr_str}", bullet_style))
                    if edu.details:
                        elements.append(Paragraph(f"{edu.details}", sub_bullet_style))
                elements.append(Spacer(1, 3))

        def add_pdf_certs():
            if hasattr(profile, 'certifications') and profile.certifications:
                add_pdf_heading("Certifications & Continuous Education", upper=(tmpl.id != "modern"))
                for cert in profile.certifications:
                    c_name = cert.name
                    c_iss = f" — <a href='https://www.deeplearning.ai' color='{tmpl.primary_color}'><u>{cert.issuer}</u></a>" if cert.issuer and "deeplearning" in cert.issuer.lower() else (f" — {cert.issuer}" if cert.issuer else "")
                    c_stat = f": {cert.status} — " if cert.status and cert.details else (f": {cert.status}" if cert.status else (f": {cert.details}" if cert.details else ""))
                    c_det = cert.details if cert.status and cert.details else ""
                    cert_line = f"• <b>{c_name}</b>{c_iss}{c_stat}{c_det}"
                    elements.append(Paragraph(cert_line, bullet_style))
                elements.append(Spacer(1, 3))

        def add_pdf_skills():
            if hasattr(profile, 'categorized_skills') and any(profile.categorized_skills.values()):
                add_pdf_heading("Technical Skills", upper=True)
                for cat, sks in profile.categorized_skills.items():
                    if sks:
                        sk_line = f"• <font color='{tmpl.primary_color}'><b>{cat}:</b></font> {', '.join(sks)}"
                        elements.append(Paragraph(sk_line, bullet_style))
                elements.append(Spacer(1, 3))
            elif profile.skills:
                add_pdf_heading("Technical Skills", upper=True)
                elements.append(Paragraph(f"• <font color='{tmpl.primary_color}'><b>Core Competencies:</b></font> {', '.join(profile.skills)}", bullet_style))
                elements.append(Spacer(1, 3))

        def add_pdf_projects():
            if profile.projects:
                add_pdf_heading("Production & Practical Projects", upper=(tmpl.id != "modern"))
                for proj in profile.projects:
                    proj_name = proj.name.strip(" :—")
                    proj_name_fmt = proj_name.replace("R²", "R<sup>2</sup>")
                    tech_part = f" | {', '.join(proj.technologies)}" if proj.technologies else ""
                    repo_part = f" | <a href='{proj.repository.strip()}' color='{tmpl.primary_color}'><u>{proj.repository.strip()}</u></a>" if proj.repository else ""
                    p_title_str = f"• <b>{proj_name_fmt}</b>{tech_part}{repo_part}"
                    elements.append(Paragraph(p_title_str, bullet_style))

                    if proj.subtitle:
                        sub_t = proj.subtitle.replace("R²", "R<sup>2</sup>")
                        elements.append(Paragraph(f"<b>{sub_t}</b>", sub_bullet_style))

                    if proj.bullets:
                        for b in proj.bullets:
                            b_clean = b.replace("R²", "R<sup>2</sup>")
                            elements.append(Paragraph(f"• {b_clean}", sub_bullet_style))
                    elif proj.description:
                        desc_clean = proj.description.replace("R²", "R<sup>2</sup>")
                        elements.append(Paragraph(desc_clean, sub_bullet_style))

                    elements.append(Spacer(1, 2))
                elements.append(Spacer(1, 2))

        def add_pdf_experience():
            if profile.experience:
                add_pdf_heading("Professional Experience", upper=(tmpl.id != "modern"))
                for exp in profile.experience:
                    clean_role = exp.role
                    if clean_role.lower() in ["professional role", "work professional", "lead professional", "professional profile", "role"]:
                        clean_role = profile.headline.split("|")[0].strip() if profile.headline else "Software Engineer"
                    dur_str = f" | {exp.duration}" if exp.duration else ""
                    comp_str = f" — {exp.company}" if exp.company and exp.company not in ["Independent", "General", ""] else ""
                    r_title = f"{clean_role}{comp_str}{dur_str}"
                    elements.append(Paragraph(f"<b>{r_title}</b>", role_title_style))
                    
                    if exp.subtitle:
                        elements.append(Paragraph(f"<font color='{tmpl.primary_color}'><b>{exp.subtitle}</b></font>", sub_bullet_style))
                    
                    if exp.summary:
                        elements.append(Paragraph(exp.summary, summary_style))
                    
                    if exp.bullets:
                        for bullet in exp.bullets:
                            bullet_clean = bullet.replace("R²", "R<sup>2</sup>")
                            elements.append(Paragraph(f"• {bullet_clean}", bullet_style))
                    
                    elements.append(Spacer(1, 2))
                elements.append(Spacer(1, 2))

        # Flow Sequencing per Template
        if tmpl.id == "corporate_elite":
            if profile.summary:
                elements.append(Paragraph(profile.summary, summary_style))
                elements.append(Spacer(1, 4))
            add_pdf_skills()
            add_pdf_projects()
            add_pdf_experience()
            add_pdf_certs()
            add_pdf_education()

        elif tmpl.id == "tech_specialist":
            if profile.summary:
                elements.append(Paragraph(profile.summary, summary_style))
                elements.append(Spacer(1, 4))
            add_pdf_skills()
            add_pdf_projects()
            add_pdf_experience()
            add_pdf_certs()
            add_pdf_education()

        elif tmpl.id == "harvard_consulting":
            if profile.summary:
                elements.append(Paragraph(profile.summary, summary_style))
                elements.append(Spacer(1, 4))
            add_pdf_experience()
            add_pdf_education()
            add_pdf_certs()
            add_pdf_skills()
            add_pdf_projects()

        else: # modern
            if profile.summary:
                elements.append(Paragraph(profile.summary, summary_style))
                elements.append(Spacer(1, 4))
            add_pdf_education()
            add_pdf_certs()
            add_pdf_skills()
            add_pdf_projects()
            add_pdf_experience()

        if profile.additional_background:
            add_pdf_heading("ADDITIONAL BACKGROUND :", upper=True)
            elements.append(Paragraph(profile.additional_background, summary_style))
            elements.append(Spacer(1, 3))

        try:
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.build(elements, canvasmaker=NumberedCanvas)
            return output_path
        except (OSError, PermissionError) as pe:
            import tempfile
            temp_out = str(Path(tempfile.gettempdir()) / Path(output_path).name)
            doc_temp = SimpleDocTemplate(temp_out, pagesize=letter, leftMargin=30, rightMargin=30, topMargin=25, bottomMargin=25)
            doc_temp.build(elements, canvasmaker=NumberedCanvas)
            return temp_out


    @classmethod
    def export_tailored_documents(
        cls,
        profile: UserProfile,
        job_title: str = "",
        company: str = "",
        original_filename: Optional[str] = None,
        template_id: str = "modern",
    ) -> Tuple[str, str]:
        """
        Generates both .docx and .pdf files with short, clean, recruiter-friendly filenames and selected template styling.
        Example: 'Mudather_Mohammed_Resume.pdf' or 'Mudather_Mohammed_Resume.docx'
        """
        if original_filename:
            raw_base = Path(original_filename).name
            raw_base = re.sub(r'(\.pdf|\.docx|\.doc|\.txt|\.cv|_cv|-cv)+', '', raw_base, flags=re.I)
            clean_name = re.sub(r"[^\w\- ]", "", raw_base).strip().replace(" ", "_")
            if not clean_name.lower().endswith("resume"):
                clean_name = f"{clean_name}_Resume"
            parts = [p.capitalize() if p.islower() or p.isupper() else p for p in clean_name.split("_") if p]
            base_name = "_".join(parts)
        else:
            candidate_clean = "_".join([w.capitalize() for w in profile.full_name.split() if w.strip()]) or "Candidate"
            base_name = f"{candidate_clean}_Resume"
        base_name = cls._sanitize_filename(base_name)
        import uuid
        uid = uuid.uuid4().hex[:8]
        docx_path = str(OUTPUT_DIR / f"{base_name}_{template_id}_{uid}.docx")
        pdf_path = str(OUTPUT_DIR / f"{base_name}_{template_id}_{uid}.pdf")

        cls.generate_docx(profile, docx_path, template_id=template_id)
        cls.generate_pdf(profile, pdf_path, template_id=template_id)

        return docx_path, pdf_path

    @classmethod
    def get_starter_profile(cls, template_id: str = "modern") -> UserProfile:
        """
        Returns a clean, beautifully formatted starter profile with clear XYZ bullet guidance
        for candidates to download, edit, and upload.
        """
        from core.resume_parser import ContactInfo, Project, WorkExperience, Education, Certification

        if template_id == "harvard_consulting":
            # 1. Official Harvard Mignone Center (MCS) & MBB Strategy Standard
            name = "[FIRST NAME] [LAST NAME]"
            headline = "Candidate for Strategy & Management Consulting | Operations & Analytics"
            summary = "High-impact, analytical professional with proven expertise in quantitative data modeling, cross-functional leadership, and strategic decision-making. Track record delivering measurable efficiency and multi-million-dollar growth for global clients."
            skills = {
                "Analytical & Financial Tools": ["Financial Modeling", "Advanced Excel / VBA", "SQL", "Tableau", "Power BI", "Python (Pandas)"],
                "Strategic Competencies": ["Quantitative Analysis", "Market Sizing & Due Diligence", "Competitive Benchmarking", "Executive Presentations", "P&L Management"],
                "Languages & Professional Skills": ["English (Fluent / Professional)", "Arabic (Native)", "Cross-Functional Leadership", "Stakeholder Management"]
            }
            contact = ContactInfo(
                email="your.email@alumni.harvard.edu",
                phone="+1 (555) 012-3456",
                location="Boston, MA (Open to Global Relocation & Remote)",
                linkedin="https://linkedin.com/in/yourprofile",
                github="https://github.com/yourhandle"
            )
            projects = [
                Project(
                    name="Global Market Expansion Due Diligence",
                    subtitle="Strategic Growth Strategy Initiative",
                    technologies=["Excel / VBA", "Financial Modeling", "Market Sizing", "Tableau"],
                    repository="https://github.com/yourhandle/consulting-market-model",
                    bullets=[
                        "Spearheaded comprehensive market sizing study across 5 international territories, identifying $4.2M in annual revenue opportunities through quantitative regression analysis.",
                        "Formulated dynamic scenario stress-testing model in Excel/VBA, evaluating downside capital risks across 12 macroeconomic conditions.",
                        "Synthesized findings into a 20-page executive presentation delivered directly to senior leadership, securing immediate board consensus."
                    ]
                )
            ]
            experience = [
                WorkExperience(
                    company="Global Management Advisory Group",
                    role="Management Strategy Consultant",
                    duration="Month Year – Present",
                    location="New York, NY / Remote",
                    subtitle="Enterprise Transformation & Strategy Practice",
                    bullets=[
                        "Led a cross-functional workstream of 8 client stakeholders to redesign procurement operations, reducing turnaround cycle time by 35%.",
                        "Constructed cost-optimization framework analyzing $50M+ in operational expenditures, capturing $3.8M in annualized savings.",
                        "Coached and mentored 4 junior analysts on hypothesis-driven problem solving and structured executive communication."
                    ]
                )
            ]
            certifications = [
                Certification(
                    name="Certified Management Consultant (CMC)",
                    issuer="Institute of Management Consultants",
                    status="Completed",
                    details="Executive Strategy, Organizational Design, and Client Governance."
                )
            ]
            education = [
                Education(
                    institution="Harvard University / College of Business",
                    degree="Bachelor of Arts / Science in Economics & Computer Science",
                    year="Graduation: Month Year | GPA: 3.8 / 4.0",
                    details="Honors: Dean's List (All Semesters). Coursework: Microeconomics, Quantitative Methods, Corporate Finance, Strategic Management."
                )
            ]

        elif template_id == "tech_specialist":
            # 2. Official Stanford BEAM & Silicon Valley Tech & AI Standard
            name = "[FIRST NAME] [LAST NAME]"
            headline = "Software & AI Engineer | Distributed Systems & Machine Learning"
            summary = "Results-driven Software & AI Engineer with expertise architecting low-latency backend microservices, high-throughput machine learning pipelines, and autonomous AI agents. Passionate about resilient systems and open-source software."
            skills = {
                "Languages & Core": ["Python", "C++", "SQL", "TypeScript", "Bash / Linux", "Go"],
                "Machine Learning & AI": ["PyTorch", "TensorFlow", "Scikit-learn", "Transformers", "LLM APIs", "Prompt Engineering", "RAG Systems"],
                "Backend & Distributed Systems": ["FastAPI", "gRPC", "Docker", "Kubernetes", "Redis", "Celery", "RESTful APIs"],
                "Databases & Storage": ["PostgreSQL", "Supabase", "Vector Databases (ChromaDB / Pinecone / pgvector)", "Redis"],
                "Cloud & DevOps": ["AWS (EC2, S3, Lambda)", "GCP", "GitHub Actions CI/CD", "Linux CLI", "vLLM Inference"]
            }
            contact = ContactInfo(
                email="your.email@stanford.edu",
                phone="+1 (555) 012-3456",
                location="San Francisco, CA / Worldwide Remote",
                linkedin="https://linkedin.com/in/yourprofile",
                github="https://github.com/yourhandle"
            )
            projects = [
                Project(
                    name="Autonomous Multi-Agent Workflow Orchestrator",
                    subtitle="Distributed Agentic Pipeline & Real-Time Queue",
                    technologies=["Python", "FastAPI", "PyTorch", "Redis", "Docker"],
                    repository="https://github.com/yourhandle/autonomous-ai-agent",
                    bullets=[
                        "Architected asynchronous event-driven workflow engine processing 10,000+ simulated jobs daily with sub-250ms latency.",
                        "Implemented in-memory Redis caching layer, reducing database compute load by 45% across distributed worker nodes.",
                        "Containerized multi-service architecture with Docker Compose and automated CI/CD deployment via GitHub Actions."
                    ]
                ),
                Project(
                    name="Enterprise Semantic Search & RAG Pipeline",
                    subtitle="High-Precision Vector Retrieval System",
                    technologies=["Python", "LangChain", "ChromaDB", "FastAPI", "OpenAI API"],
                    repository="https://github.com/yourhandle/semantic-search-rag",
                    bullets=[
                        "Indexed 500,000+ domain documents with sub-100ms vector lookup, achieving 94% factual answering precision.",
                        "Integrated cross-encoder reranking algorithms, eliminating hallucinations across multi-turn reasoning workflows."
                    ]
                )
            ]
            experience = [
                WorkExperience(
                    company="Silicon Valley Engineering Labs",
                    role="Senior Software & AI Engineer",
                    duration="Month Year – Present",
                    location="Palo Alto, CA / Remote",
                    subtitle="Core Infrastructure & Machine Learning Team",
                    bullets=[
                        "Engineered distributed data ingestion pipelines handling 50M+ monthly events with 99.9% uptime.",
                        "Integrated real-time model inference endpoints in FastAPI, improving system throughput by 3.2x under high concurrency.",
                        "Collaborated in an Agile squad of 8 engineers, conducting rigorous code reviews and maintaining 92% test coverage."
                    ]
                )
            ]
            certifications = [
                Certification(
                    name="Deep Learning Specialization",
                    issuer="DeepLearning.AI",
                    status="Completed",
                    details="Neural Networks, CNNs, Sequence Models, and Hyperparameter Optimization."
                ),
                Certification(
                    name="AWS Certified Solutions Architect – Associate",
                    issuer="Amazon Web Services",
                    status="Completed",
                    details="High-Availability Cloud Architecture, Distributed Storage, and Serverless Infrastructure."
                )
            ]
            education = [
                Education(
                    institution="Stanford University / Institute of Technology",
                    degree="Bachelor of Science in Computer Science & Artificial Intelligence",
                    year="2021 – 2025",
                    details="Relevant Coursework: Algorithms & Data Structures, Distributed Systems, Deep Learning, Operating Systems."
                )
            ]

        elif template_id == "corporate_elite":
            # 3. Official Wharton & Wall Street Oasis (WSO) Executive Standard
            name = "[FIRST NAME] [LAST NAME]"
            headline = "Corporate Executive & Senior Technology Leader"
            summary = "Strategic, performance-focused technology executive with extensive track record directing high-scale engineering programs, modernizing enterprise architectures, and generating multi-million-dollar operational efficiencies."
            skills = {
                "Executive Leadership": ["Engineering Management", "P&L Accountability", "Cross-Functional Leadership", "Stakeholder Alignment"],
                "Enterprise Architecture": ["Distributed Systems", "Cloud Transformation (AWS/GCP)", "Enterprise AI Integration", "REST/Microservices"],
                "Governance & Delivery": ["Agile/Scrum Governance", "Data Privacy & Security", "Risk Management", "Vendor Negotiations"],
                "Data & Engineering Tools": ["Python", "SQL", "FastAPI", "Docker", "Supabase/PostgreSQL", "Power BI"]
            }
            contact = ContactInfo(
                email="your.email@wharton.upenn.edu",
                phone="+1 (555) 012-3456",
                location="New York, NY / Remote",
                linkedin="https://linkedin.com/in/yourprofile",
                github="https://github.com/yourhandle"
            )
            projects = [
                Project(
                    name="Enterprise Digital Modernization Initiative",
                    subtitle="Core Infrastructure Transformation Program",
                    technologies=["Cloud Architecture", "FastAPI", "PostgreSQL", "Docker"],
                    repository="https://github.com/yourhandle/enterprise-transformation",
                    bullets=[
                        "Spearheaded modernization of legacy transaction pipelines, boosting throughput capacity from 10M to 100M daily records.",
                        "Negotiated primary cloud vendor licensing agreements, achieving 30% reduction in annual compute expenditure ($850K savings)."
                    ]
                )
            ]
            experience = [
                WorkExperience(
                    company="Fortune 500 Enterprise Technologies",
                    role="Senior Director of Engineering",
                    duration="Month Year – Present",
                    location="New York, NY",
                    subtitle="Global Platform & Enterprise Solutions Squad",
                    bullets=[
                        "Directed 18+ senior engineers across 3 squads, delivering flagship enterprise platform on-time and $350K under budget.",
                        "Architected scalable cloud microservices migration, reducing annual infrastructure expenditures by 28% ($1.2M savings).",
                        "Instituted secure coding practices and automated testing frameworks, cutting critical production incidents by 75%."
                    ]
                )
            ]
            certifications = [
                Certification(
                    name="Project Management Professional (PMP)",
                    issuer="Project Management Institute (PMI)",
                    status="Completed",
                    details="Executive Program Management, Budget Governance, and Risk Mitigation."
                )
            ]
            education = [
                Education(
                    institution="The Wharton School / University of Pennsylvania",
                    degree="Bachelor of Science in Economics / Engineering Management",
                    year="Graduated with Honors",
                    details="Executive Leadership Fellow. Core focus: Corporate Finance, Technology Innovation, Enterprise Operations."
                )
            ]

        else: # modern (Google FAANG "XYZ Formula" Engineering Standard)
            # 4. Official Google FAANG "Accomplished [X] as measured by [Y], by doing [Z]" Standard
            name = "[FIRST NAME] [LAST NAME]"
            headline = "Senior Software Engineer | High-Scale Distributed Platforms"
            summary = "Results-driven Software Engineer with 5+ years of experience delivering high-availability backend systems and machine learning services using outcome-driven engineering principles."
            skills = {
                "Core Programming Languages": ["Python", "Go", "C++", "SQL", "TypeScript"],
                "Distributed Systems & APIs": ["FastAPI", "gRPC", "Microservices", "Event-Driven Queues", "Redis"],
                "Cloud & Container Platforms": ["Google Cloud Platform (GCP)", "Kubernetes (K8s)", "Docker", "CI/CD Pipelines"],
                "AI & Machine Learning": ["PyTorch", "TensorFlow", "Scikit-learn", "Vector Search", "LLM APIs"]
            }
            contact = ContactInfo(
                email="your.email@gmail.com",
                phone="+1 (555) 012-3456",
                location="San Francisco, CA / Remote",
                linkedin="https://linkedin.com/in/yourprofile",
                github="https://github.com/yourhandle"
            )
            projects = [
                Project(
                    name="High-Performance Distributed Data Indexer",
                    subtitle="Real-Time Vector Embedding Service",
                    technologies=["Python", "FastAPI", "Docker", "PostgreSQL"],
                    repository="https://github.com/yourhandle/distributed-indexer",
                    bullets=[
                        "Accomplished sub-second indexing across 1,000,000+ vector records (as measured by load testing metrics) by developing parallelized batch worker pools.",
                        "Accomplished 95% automated test coverage across core libraries by establishing end-to-end integration pipelines with GitHub Actions."
                    ]
                )
            ]
            experience = [
                WorkExperience(
                    company="Global Technology Platforms",
                    role="Senior Backend Software Engineer",
                    duration="Month Year – Present",
                    location="Mountain View, CA / Remote",
                    subtitle="Core Systems & Distributed Reliability Group",
                    bullets=[
                        "Accomplished 40% reduction in API response latency (from 180ms to 108ms) by redesigning backend data serialization in FastAPI and implementing Redis memory caching.",
                        "Accomplished $450,000 annual cloud infrastructure cost savings by optimizing container orchestration and autoscaling policies across Kubernetes clusters.",
                        "Accomplished 99.95% system uptime across 12 distributed microservices by architecting automated failover queues and circuit-breaker patterns."
                    ]
                )
            ]
            certifications = [
                Certification(
                    name="Google Cloud Certified Professional Cloud Architect",
                    issuer="Google Cloud",
                    status="Completed",
                    details="Scalable Cloud Architecture, Kubernetes Cluster Management, and High-Reliability Infrastructure."
                )
            ]
            education = [
                Education(
                    institution="University of California, Berkeley",
                    degree="Bachelor of Science in Computer Science",
                    year="2018 – 2022",
                    details="Graduated with Honors. Core coursework: Distributed Systems, Operating Systems, Database Architecture, Machine Learning."
                )
            ]

        flat_skills = [s for sub in skills.values() for s in sub]

        contact = ContactInfo(
            email="your.email@example.com",
            phone="+1 (555) 012-3456",
            location="City, State / Worldwide Remote",
            linkedin="https://linkedin.com/in/yourprofile",
            github="https://github.com/yourhandle"
        )

        projects = [
            Project(
                name="Autonomous Multi-Agent Orchestrator",
                subtitle="Distributed AI Workflow Pipeline",
                technologies=["Python", "FastAPI", "PyTorch", "Redis", "Docker"],
                repository="https://github.com/yourhandle/autonomous-ai-agent",
                bullets=[
                    "Accomplished sub-second execution latency, as measured by 10,000+ simulated jobs daily, by designing async event-driven queues.",
                    "Engineered custom memory cache with Redis, cutting RAM consumption by 45% across containerized workers."
                ]
            ),
            Project(
                name="Enterprise Semantic Search & RAG Engine",
                subtitle="High-Precision Retrieval System",
                technologies=["Python", "LangChain", "ChromaDB", "FastAPI", "OpenAI API"],
                repository="https://github.com/yourhandle/semantic-search-rag",
                bullets=[
                    "Indexed 500,000+ domain documents with sub-100ms vector lookup, improving factual answering accuracy to 94%.",
                    "Integrated cross-encoder reranking algorithms, eliminating hallucinations across multi-turn reasoning queries."
                ]
            )
        ]

        experience = [
            WorkExperience(
                company="Global Enterprise Technologies",
                role="Senior AI Engineer",
                duration="2021 – Present",
                location="San Francisco, CA / Remote",
                subtitle="Core Machine Learning & Infrastructure Squad",
                bullets=[
                    "Spearheaded enterprise AI agent deployment across 12 services, reducing operational cycle time by 60%.",
                    "Architected high-throughput data processing pipelines handling 50M+ requests monthly with 99.9% uptime.",
                    "Mentored junior engineers and instituted automated testing and linting standards across the codebase."
                ]
            )
        ]

        certifications = [
            Certification(
                name="Deep Learning Specialization",
                issuer="DeepLearning.AI",
                status="Completed",
                details="Neural Networks, CNNs, Sequence Models, and Hyperparameter Optimization."
            ),
            Certification(
                name="AWS Certified Solutions Architect",
                issuer="Amazon Web Services",
                status="Completed",
                details="Cloud Architecture, High-Availability Infrastructure, and Serverless Systems."
            )
        ]

        education = [
            Education(
                institution="University / Institute of Technology",
                degree="Bachelor of Science in Computer Science / Engineering",
                year="2017 – 2021",
                details="Graduated with Honors. Core coursework: Distributed Systems, Algorithms, Machine Learning, Operating Systems."
            )
        ]

        return UserProfile(
            full_name=name,
            headline=headline,
            contact=contact,
            summary=summary,
            skills=flat_skills,
            categorized_skills=skills,
            projects=projects,
            experience=experience,
            certifications=certifications,
            education=education
        )

    @classmethod
    def generate_starter_template(cls, template_id: str, output_path: str, format: str = "docx") -> str:
        """
        Generates a clean starter template in DOCX or PDF format for the specified design style.
        """
        starter_prof = cls.get_starter_profile(template_id=template_id)
        if format.lower() == "pdf":
            return cls.generate_pdf(starter_prof, output_path, template_id=template_id)
        return cls.generate_docx(starter_prof, output_path, template_id=template_id)
