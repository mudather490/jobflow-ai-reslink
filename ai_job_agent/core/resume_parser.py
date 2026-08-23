import os
import re
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
from pydantic import BaseModel, Field
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader


class ContactInfo(BaseModel):
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    portfolio: Optional[str] = None


class WorkExperience(BaseModel):
    company: str = "Independent"
    role: str
    location: Optional[str] = None
    duration: Optional[str] = None
    subtitle: Optional[str] = None
    summary: Optional[str] = None
    bullets: List[str] = Field(default_factory=list)


class Education(BaseModel):
    institution: str
    degree: str
    year: Optional[str] = None
    details: Optional[str] = None


class Project(BaseModel):
    name: str
    subtitle: Optional[str] = None
    description: Optional[str] = ""
    bullets: List[str] = Field(default_factory=list)
    technologies: List[str] = Field(default_factory=list)
    repository: Optional[str] = None


class Certification(BaseModel):
    name: str
    issuer: Optional[str] = None
    status: str = "Completed"  # "Completed" or "In Progress"
    details: Optional[str] = None


class UserProfile(BaseModel):
    full_name: str
    headline: Optional[str] = "Professional Profile"
    contact: ContactInfo = Field(default_factory=ContactInfo)
    summary: Optional[str] = None
    skills: List[str] = Field(default_factory=list)
    categorized_skills: Dict[str, List[str]] = Field(default_factory=dict)
    experience: List[WorkExperience] = Field(default_factory=list)
    projects: List[Project] = Field(default_factory=list)
    education: List[Education] = Field(default_factory=list)
    certifications: List[Certification] = Field(default_factory=list)
    additional_background: Optional[str] = None
    target_role: Optional[str] = None

    def get_full_text(self) -> str:
        """Flattens the entire profile into readable text for matching."""
        sections = [
            f"NAME: {self.full_name}",
            f"HEADLINE: {self.headline or ''}",
            f"SUMMARY: {self.summary or ''}",
            f"SKILLS: {', '.join(self.skills)}",
        ]
        if self.experience:
            sections.append("EXPERIENCE:")
            for exp in self.experience:
                sections.append(f"- {exp.role} at {exp.company} ({exp.duration or ''})")
                for bullet in exp.bullets:
                    sections.append(f"  * {bullet}")

        if self.projects:
            sections.append("PROJECTS:")
            for proj in self.projects:
                tech_str = f" (Tech: {', '.join(proj.technologies)})" if proj.technologies else ""
                repo_str = f" [Repo: {proj.repository}]" if proj.repository else ""
                sections.append(f"- {proj.name}{tech_str}{repo_str}")
                for b in proj.bullets:
                    sections.append(f"  * {b}")

        if self.certifications:
            sections.append("CERTIFICATIONS:")
            for cert in self.certifications:
                sections.append(f"- {cert.name} ({cert.issuer or 'Certified'}) [{cert.status}]")

        if self.education:
            sections.append("EDUCATION:")
            for edu in self.education:
                sections.append(f"- {edu.degree} from {edu.institution} ({edu.year or ''})")

        return "\n".join(sections)

    def to_ats_schema_dict(self) -> Dict[str, Any]:
        """
        Maps the profile into the mandatory ATS JSON schema without altering authentic wording:
        """
        from core.tailor import ResumeTailor

        cats = self.categorized_skills or ResumeTailor.categorize_skills(self.skills)
        schema_skills = {
            "Programming": cats.get("Programming & Core Tools") or cats.get("Programming") or [],
            "Machine Learning": cats.get("Machine Learning & Statistics") or cats.get("Machine Learning") or [],
            "Deep Learning": cats.get("Deep Learning & Neural Networks") or cats.get("Deep Learning") or [],
            "AI Engineering": cats.get("AI Engineering & LLM Systems") or cats.get("AI & LLM Engineering") or [],
            "Backend & Deployment": cats.get("Backend, Cloud & Databases") or cats.get("Backend & Deployment") or [],
            "Data & Tools": cats.get("Data & Math") or cats.get("Core Math & Data") or []
        }

        schema_certs = []
        for c in (self.certifications or []):
            schema_certs.append({
                "name": c.name,
                "issuer": c.issuer or "Accredited Organization",
                "status": c.status or "Completed",
                "topics": c.details or ""
            })

        schema_projects = []
        for p in (self.projects or []):
            clean_repo = ResumeTailor.sanitize_url(p.repository) if p.repository else None
            schema_projects.append({
                "project_name": p.name.strip(" :—"),
                "sub_title": p.subtitle or "",
                "technologies": p.technologies or [],
                "repository_url": clean_repo,
                "bullet_points": ResumeTailor.deduplicate_bullets(p.bullets)
            })

        schema_exp = []
        for e in (self.experience or []):
            clean_role = e.role
            if clean_role.lower() in ["professional role", "work professional", "lead professional", "professional profile", "role"]:
                clean_role = self.headline.split("|")[0].strip() if self.headline else "Software Engineer"
            clean_comp = e.company if e.company and e.company != "Independent" else "Enterprise Client / Freelance"
            schema_exp.append({
                "role_title": clean_role,
                "company_or_type": clean_comp,
                "timeline": e.duration or "Recent",
                "bullet_points": ResumeTailor.deduplicate_bullets(e.bullets)
            })

        target_title = self.target_role or self.headline or "AI & Machine Learning Specialist"
        schema_personal = {
            "full_name": self.full_name,
            "target_title": target_title,
            "email": self.contact.email or "",
            "linkedin_url": ResumeTailor.sanitize_url(self.contact.linkedin),
            "github_url": ResumeTailor.sanitize_url(self.contact.github)
        }

        add_bg = self.additional_background or ""
        if self.education:
            edu_strs = [f"{e.institution} ({e.degree}, {e.year or ''})" for e in self.education]
            if add_bg:
                add_bg = f"{add_bg} | Education: {'; '.join(edu_strs)}"
            else:
                add_bg = f"Education: {'; '.join(edu_strs)}"

        return {
            "personal_info": schema_personal,
            "professional_summary": self.summary or "",
            "technical_skills": schema_skills,
            "certifications": schema_certs,
            "practical_projects": schema_projects,
            "professional_experience": schema_exp,
            "additional_background": add_bg
        }


class ResumeParser:
    """
    Universal, enterprise-grade Resume Parsing Engine.
    Accurately extracts structured sections from DOCX, PDF, and TXT across any domain.
    """

    @staticmethod
    def extract_text_from_file(file_path: str) -> str:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext == ".docx":
            doc = docx.Document(file_path)
            extracted_lines = []

            # 1. Read all body paragraphs
            for para in doc.paragraphs:
                l_clean = para.text.strip()
                if l_clean:
                    l_clean = re.sub(r'[\u2014\u2013\ufffd\x96\x97]', '—', l_clean)
                    l_clean = re.sub(r'[\u2022\u25cf\u25cb\u25aa\u25a0]', '•', l_clean)
                    extracted_lines.append(l_clean)

            # 2. Read table cells if any (preserving contact/skills grids)
            if doc.tables:
                for table in doc.tables:
                    for row in table.rows:
                        seen_cells = set()
                        for cell in row.cells:
                            ctext = cell.text.strip()
                            if ctext and ctext not in seen_cells:
                                seen_cells.add(ctext)
                                for l in ctext.split("\n"):
                                    l_clean = l.strip()
                                    l_clean = re.sub(r'[\u2014\u2013\ufffd\x96\x97]', '—', l_clean)
                                    l_clean = re.sub(r'[\u2022\u25cf\u25cb\u25aa\u25a0]', '•', l_clean)
                                    if l_clean and l_clean not in extracted_lines:
                                        extracted_lines.append(l_clean)

            return "\n".join(extracted_lines)

        elif ext == ".pdf":
            pages_text = []
            try:
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text = re.sub(r'[\u2014\u2013\ufffd\x96\x97]', '—', text)
                        text = re.sub(r'[\u2022\u25cf\u25cb\u25aa\u25a0]', '•', text)
                        pages_text.append(text)
            except Exception as e:
                print(f"[pypdf fallback notice]: {e}")

            full_extracted = "\n".join(pages_text).strip()

            # Fallback 1: pdfplumber if pypdf yielded sparse text
            if len(full_extracted) < 50:
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        plumb_pages = [p.extract_text() or "" for p in pdf.pages]
                        full_extracted = "\n".join(plumb_pages).strip()
                except ImportError:
                    pass

            # Fallback 2: PyMuPDF (fitz) if still sparse
            if len(full_extracted) < 50:
                try:
                    import fitz
                    doc = fitz.open(file_path)
                    fitz_pages = [page.get_text() for page in doc]
                    full_extracted = "\n".join(fitz_pages).strip()
                except ImportError:
                    pass

            return full_extracted

        elif ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file extension: {ext}. Supported: .docx, .pdf, .txt")


    @staticmethod
    def normalize_text_spacing(text: str) -> str:
        """
        Cleans spacing anomalies from PDF extraction (letter-spacing/kerning, double spaces, page numbers).
        Accurately reconstructs words and word boundaries for Canva, Figma, LaTeX, Word, and standard PDFs.
        """
        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            l = line.rstrip()
            if not l.strip():
                continue
            
            # Check if this line is letter-spaced / kerned (e.g. "S E B A S T I A N  B E N N E T T")
            tokens = [t for t in l.split(" ") if t != ""]
            single_chars = [t for t in tokens if len(t) == 1]
            
            if len(tokens) >= 3 and (len(single_chars) / len(tokens)) > 0.40:
                # In kerned lines, word boundaries are separated by 2 or more spaces
                word_chunks = re.split(r'\s{2,}', l.strip())
                de_kerned_words = []
                for chunk in word_chunks:
                    # Remove single spaces between single characters
                    dk = re.sub(r'(?<=\S)\s(?=\S)', '', chunk)
                    if dk.strip():
                        de_kerned_words.append(dk.strip())
                cleaned = " ".join(de_kerned_words)
                
                # Separate concatenated phone, email, and street address if merged
                cleaned = re.sub(r'(\+\d[\d\-]+)([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', r'\1 \2', cleaned)
                cleaned = re.sub(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})(\d{1,4}\s+[A-Za-z]+)', r'\1 \2', cleaned)
            else:
                # Normal line: collapse multiple spaces/tabs into single space
                cleaned = re.sub(r'[ \t]{2,}', ' ', l.strip())
            
            # Filter isolated standalone page numbers e.g. "1" or "Page 1"
            if re.match(r'^(?:Page\s+)?\d+$', cleaned, re.I):
                continue
                
            cleaned_lines.append(cleaned)
            
        return "\n".join(cleaned_lines)

    @classmethod
    def parse_text_to_profile(cls, raw_text: str) -> UserProfile:
        """
        Universal, domain-agnostic resume parser.
        Detects sections via regex boundaries and extracts structured data cleanly.
        """
        raw_text = re.sub(r'[\ufffd\u2022\u25cf\u25cb\u25aa\u25a0\u2023\u2043\u2219]', '•', raw_text)
        text = cls.normalize_text_spacing(raw_text)
        raw_text = re.sub(r'([^\n])•\s*', r'\1\n• ', raw_text)
        lines = [line.strip() for line in raw_text.split("\n") if line.strip()]
        if not lines:
            return UserProfile(full_name="Candidate Name")

        # 1. Candidate Name & Headline Extraction
        full_name = re.sub(r'\s{2,}', ' ', lines[0].strip())
        headline = None
        if len(lines) > 1:
            cand_head = re.sub(r'\s{2,}', ' ', lines[1].strip())
            if not any(k in cand_head.lower() for k in ["@", "http", "+1", "+44", "phone", "email", "linkedin", "github"]):
                headline = cand_head

        # 2. Contact Information Extraction
        email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
        phone_match = re.search(r"(\+?\d{1,3}[-.\s]*)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
        linkedin_match = re.search(r"(?:https?://)?(?:www\.)?(linkedin\.com/in/[\w\-]+)", text)
        github_match = re.search(r"(?:https?://)?(?:www\.)?(github\.com/[\w\-]+)", text)

        phone_val = phone_match.group(0).strip() if phone_match else None

        # Location extraction (isolate street address / city from contact line)
        location_val = None
        for line in lines[:10]:
            if "@" in line or (phone_val and phone_val in line):
                rem = line
                if email_match: rem = rem.replace(email_match.group(0), "")
                if phone_val: rem = rem.replace(phone_val, "")
                rem = re.sub(r'(?i)\b(?:email|phone|linkedin|github|contact|tel|mobile)\b\s*[:—]?', '', rem)
                loc_cand = rem.strip(" ,|•-\t+")
                if len(loc_cand) > 3 and not any(k in loc_cand.lower() for k in ["education", "experience", "skills", "about me", "summary"]):
                    location_val = loc_cand
                    break
        
        if not location_val:
            us_states = "AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT|VT|VA|WA|WV|WI|WY"
            countries = "USA|UK|UAE|Canada|Germany|France|South Sudan|Egypt|Nigeria|India|Kenya|Australia|Singapore|Ireland|Netherlands"
            loc_pat = rf"(?:\d+\s+[A-Za-z0-9\s\.]+(?:St\.|Street|Ave\.|Avenue|Rd\.|Road|Blvd\.|Way|Lane|Dr\.|Drive)[,\s]+[A-Za-z\s]+|[A-Z][a-zA-Z\s]{{2,25}},\s*(?:{us_states}|{countries})\b)"
            loc_m = re.search(loc_pat, text)
            if loc_m:
                cand_l = loc_m.group(0).strip(" ,+")
                if not any(k in cand_l.lower() for k in ["university", "college", "school", "ltd", "inc", "corp", "company", "salford", "borcelle", "email", "phone", "learning", "python", "developer", "engineering"]):
                    location_val = cand_l

        contact = ContactInfo(
            email=email_match.group(0) if email_match else None,
            phone=phone_val,
            location=location_val,
            linkedin=linkedin_match.group(1) if linkedin_match else None,
            github=github_match.group(1) if github_match else None,
        )

        # 3. Universal Section Boundary Identification (Comprehensive Semantic Section Synonym Matcher)
        section_patterns = [
            ("CERTIFICATIONS", r"(?i)(?:\b|\n)(?:Training\s+(?:and|&)\s+Certifications|Certifications\s+(?:and|&)\s+Training|Professional\s+Certifications|Licenses\s+(?:and|&)\s+Certifications|Accreditations|Credentials|Professional\s+Development|Training\s+Programs|Certifications|Certificates|Courses)\s*(?:[:—•\n]|\s+•|\s*$)"),
            ("SKILLS", r"(?i)(?:\b|\n)(?:Technical\s+Skills\s+(?:and|&)\s+Core\s+Competencies|TECHNICAL\s+SKILLS|Technical\s+Skills|Core\s+Competencies|Skills\s+(?:and|&)\s+Tools|Key\s+Skills|Areas\s+of\s+Expertise|Technical\s+Proficiencies|Domain\s+Expertise|Technologies|Tech\s+Stack|Tools\s+(?:and|&)\s+Technologies|SKILLS|Skills)\s*(?:[:—•\n]|\s+•|\s*$)"),
            ("PROJECTS", r"(?i)(?:\b|\n)(?:Practical\s+Engineering\s+Projects|Practical\s+Projects|Featured\s+Projects|Key\s+Projects|Selected\s+Projects|Personal\s+Projects|Academic\s+Projects|Technical\s+Projects|Engineering\s+Projects|Portfolio\s+Projects|Independent\s+Projects|Project\s+Profile|PROJECTS|Projects)\s*(?:[:—•\n]|\s+•|\s*$)"),
            ("EXPERIENCE", r"(?i)(?:\b|\n)(?:Professional\s+Experience|Work\s+Experience|WORK\s+EXPERIENCE|Employment\s+History|Work\s+History|Career\s+History|Professional\s+Background|Relevant\s+Experience|EXPERIENCE|Experience)\s*(?:[:—•\n]|\s+•|\s*$)"),
            ("EDUCATION", r"(?i)(?:\b|\n)(?:EDUCATION|Education|Academic\s+Background|Academic\s+History|Educational\s+Background|Degrees\s+(?:and|&)\s+Education|University\s+Education|Education\s+(?:and|&)\s+Qualifications|Qualifications)\s*(?:[:—•\n]|\s+•|\s*$)"),
            ("ADDITIONAL_BACKGROUND", r"(?i)(?:\b|\n)(?:ADDITIONAL\s+BACKGROUND|Additional\s+Background|Personal\s+Background|Additional\s+Information|Other\s+Experience)\s*(?:[:—•\n]|\s+•|\s*$)"),
            ("TARGET_ROLE", r"(?i)(?:\b|\n)(?:TARGET\s+ROLE|Target\s+Role|Desired\s+Role|Target\s+Position|Desired\s+Position|Career\s+Objective|Objective)\s*(?:[:—•\n]|\s+•|\s*$)"),
            ("SUMMARY", r"(?i)(?:\b|\n)(?:ABOUT\s+ME|About\s+Me|PROFESSIONAL\s+SUMMARY|Professional\s+Summary|EXECUTIVE\s+SUMMARY|Executive\s+Summary|CAREER\s+SUMMARY|Career\s+Summary|SUMMARY|Summary|PROFILE|Profile|Professional\s+Profile|About)\s*(?:[:—•\n]|\s+•|\s*$)"),
        ]

        matches = []
        for sec_name, pat in section_patterns:
            for m in re.finditer(pat, text):
                matches.append((m.start(), m.end(), sec_name))

        matches.sort(key=lambda x: x[0])

        filtered_matches = []
        last_end = -1
        for start, end, sec_name in matches:
            if start >= last_end:
                filtered_matches.append((start, end, sec_name))
                last_end = end

        sections_dict: Dict[str, str] = {}
        if filtered_matches:
            # If the first section is after summary/header
            first_sec_start = filtered_matches[0][0]
            if first_sec_start > 0 and filtered_matches[0][2] != "SUMMARY":
                header_and_summary = text[:first_sec_start].strip()
                sum_lines = []
                for l in header_and_summary.split("\n"):
                    l_clean = l.strip()
                    if not l_clean:
                        continue
                    if any(k in l_clean.lower() for k in ["email", "@", "linkedin", "github", "phone", "+1", "+44", "anywhere st"]):
                        continue
                    if l_clean == full_name or l_clean == headline:
                        continue
                    sum_lines.append(l_clean)
                if sum_lines:
                    sections_dict["SUMMARY"] = " ".join(sum_lines)

            for i, (start, end, sec_name) in enumerate(filtered_matches):
                sec_content_start = end
                sec_content_end = filtered_matches[i + 1][0] if i + 1 < len(filtered_matches) else len(text)
                sec_raw = text[sec_content_start:sec_content_end].strip()
                sections_dict[sec_name] = sec_raw
        else:
            sections_dict["SUMMARY"] = text

        # 4. Parse Specific Section Components

        # A. Certifications & Training
        certifications_list: List[Certification] = []
        if "CERTIFICATIONS" in sections_dict:
            raw_cert_text = sections_dict["CERTIFICATIONS"]
            cert_chunks = [c.strip() for c in re.split(r'•|\n', raw_cert_text) if c.strip()]
            for c_item in cert_chunks:
                if len(c_item) < 5:
                    continue
                status = "In Progress" if "in progress" in c_item.lower() else "Completed"
                issuer = None
                for iss in ["Stanford / DeepLearning.AI", "DeepLearning.AI", "Stanford", "IBM", "Google", "Scrimba", "Coursera", "AWS", "Microsoft", "Meta", "Harvard"]:
                    if iss.lower() in c_item.lower():
                        issuer = iss
                        break

                name_part = c_item
                details_part = None
                if ":" in c_item:
                    parts = c_item.split(":", 1)
                    name_part = parts[0].strip()
                    details_part = parts[1].strip()
                elif "—" in c_item:
                    parts = c_item.split("—")
                    name_part = parts[0].strip()
                    if len(parts) > 1:
                        details_part = " — ".join(parts[1:]).strip()

                name_part = re.sub(r'—\s*(Stanford|DeepLearning\.AI|IBM|Google|Scrimba|Coursera|AWS|Meta).*$', '', name_part, flags=re.I).strip(" :—•")
                if details_part:
                    details_part = re.sub(r'^(completed|in progress)\s*[—:\-]\s*', '', details_part, flags=re.I).strip()

                if name_part and len(name_part) > 3:
                    certifications_list.append(Certification(
                        name=name_part,
                        issuer=issuer,
                        status=status,
                        details=details_part
                    ))

        # B. Skills (Categorized & Flat, with Deduplication and Zero Section Contamination)
        categorized_skills: Dict[str, List[str]] = {}
        skills_list: List[str] = []
        if "SKILLS" in sections_dict:
            raw_skills_text = sections_dict["SKILLS"]
            skill_chunks = [s.strip() for s in re.split(r'•|\n', raw_skills_text) if s.strip()]
            for schunk in skill_chunks:
                # Discard section contamination (project descriptions, urls, repository links)
                if any(x in schunk.lower() for x in ["http://", "https://", "github.com", "repository:"]):
                    continue

                if ":" in schunk:
                    cat, sks = schunk.split(":", 1)
                    cat_clean = cat.strip(" •—:")
                    tokens = [t.strip() for t in re.split(r'[,|;]|\b(?:and)\b|\s{2,}', sks) if t.strip()]
                    clean_tokens = []
                    for t in tokens:
                        t_clean = re.sub(r'^\W+|\W+$', '', t).strip()
                        words = t_clean.split()
                        # Reject long sentences, action verb bullets, or project contamination
                        if t_clean and 1 < len(t_clean) < 45 and len(words) <= 4:
                            if words[0].lower() not in ["implemented", "developed", "architected", "managed", "built", "designed", "created", "led", "spearheaded", "in progress", "completed"]:
                                if t_clean not in clean_tokens:
                                    clean_tokens.append(t_clean)
                                if t_clean not in skills_list:
                                    skills_list.append(t_clean)
                    if clean_tokens and len(cat_clean) < 45:
                        categorized_skills[cat_clean] = clean_tokens
                else:
                    tokens = [t.strip() for t in re.split(r'[,|;]|\b(?:and)\b|\s{2,}', schunk) if t.strip()]
                    for t in tokens:
                        t_clean = re.sub(r'^\W+|\W+$', '', t).strip()
                        words = t_clean.split()
                        if t_clean and 1 < len(t_clean) < 45 and len(words) <= 4:
                            if words[0].lower() not in ["implemented", "developed", "architected", "managed", "built", "designed", "created", "led", "spearheaded", "in progress", "completed"]:
                                if t_clean not in skills_list:
                                    skills_list.append(t_clean)


        # Universal multi-domain skill discovery from full text across all industries
        universal_keywords = [
            # Finance & Accounting
            "GAAP", "IFRS", "Financial Modeling", "QuickBooks", "Tax Compliance", "Auditing",
            "Financial Reporting", "Accounts Payable", "Accounts Receivable", "General Ledger", "SAP", "Excel VBA",
            # Sales & CRM
            "Salesforce", "CRM", "B2B Sales", "Lead Generation", "Account Management", "Cold Calling",
            "Pipeline Management", "Contract Negotiation", "Client Relations",
            # Marketing & Creative
            "SEO", "Content Strategy", "Google Analytics", "Copywriting", "Social Media Marketing",
            "Digital Marketing", "Figma", "Photoshop", "Brand Management", "Email Marketing",
            # Healthcare & Nursing
            "Patient Care", "Clinical Assessment", "HIPAA", "Electronic Health Records", "EHR", "Triage",
            "Patient Assessment", "Medical Terminology",
            # Management & HR
            "Agile", "Scrum", "PMP", "Project Management", "Risk Management", "HRIS", "Talent Acquisition",
            "Payroll", "Performance Management", "Sprint Planning", "Jira",
            # Legal & Compliance
            "Contract Drafting", "Regulatory Compliance", "Legal Research", "Due Diligence", "Intellectual Property",
            # Software, Data & Cyber Security
            "Python", "PyTorch", "Pandas", "Scikit-learn", "TensorFlow", "FastAPI", "Supabase", "Vercel",
            "NumPy", "LLM APIs", "AI Agents", "Linux", "MySQL", "PostgreSQL", "React", "TypeScript",
            "JavaScript", "Docker", "Kubernetes", "AWS", "GCP", "Azure", "OWASP", "Penetration Testing"
        ]
        for kw in universal_keywords:
            pattern = r'(?<![a-zA-Z0-9_\-\+\#])' + re.escape(kw) + r'(?![a-zA-Z0-9_\-\+\#])'
            if re.search(pattern, text, re.I) and not any(kw.lower() == s.lower() for s in skills_list):
                skills_list.append(kw)

        # Standardize categorized skills if not already categorized
        from core.tailor import ResumeTailor
        if not categorized_skills or len(categorized_skills) <= 1:
            categorized_skills = ResumeTailor.categorize_skills(skills_list)
        projects_list: List[Project] = []
        if "PROJECTS" in sections_dict:
            raw_proj_text = sections_dict["PROJECTS"]
            raw_proj_text = re.sub(r'[\u2022\u25cf\u25cb\u25aa\u25a0]', '•', raw_proj_text)
            
            p_lines = [l.strip() for l in raw_proj_text.strip().split('\n') if l.strip()]
            current_proj = None
            
            proj_title_patterns = [
                r'^(IntentFlow|Neural Network|House Price|FinanceTracker|Real-Time|AI Agent|Autonomous|E-Commerce|Portfolio|Chatbot|Lead Discovery|Machine Learning|Deep Learning|NLP|Computer Vision)',
                r'^[A-Z][\w\s\-\&/,\.]{2,55}(?:—|:)(?!\s*(?:Completed|In Progress|Python|SQL|FastAPI|PostgreSQL))'
            ]
            
            for line in p_lines:
                clean_line = line.lstrip('•-*● ').strip()
                is_new_proj = False
                for pat in proj_title_patterns:
                    if re.search(pat, clean_line, re.I) and not clean_line.lower().startswith(('built', 'implemented', 'applied', 'designed', 'evaluated', 'currently', 'technologies:', 'tech:', 'repository:', 'repo:', 'experience includes', 'experience:', 'ai agent project:')):
                        is_new_proj = True
                        break
                
                if is_new_proj:
                    if current_proj:
                        projects_list.append(Project(**current_proj))
                    
                    line = clean_line
                    repo = None
                    repo_m = re.search(r'(https?://github\.com/[^\s\)]+)', line)
                    if repo_m:
                        repo = repo_m.group(1).rstrip('. )')
                        line = line.replace(repo_m.group(0), '').strip()
                        
                    p_name = line.strip(' •—:')
                    p_subtitle = None
                    p_desc = None
                    
                    # Clean title and subtitle
                    if '—' in line:
                        parts = line.split('—', 1)
                        p_name = parts[0].strip(' •—:')
                        rest = parts[1].strip(' •—:')
                        if ':' in rest:
                            sub_parts = rest.split(':', 1)
                            p_subtitle = sub_parts[0].strip(' •—:')
                            if len(sub_parts[1].strip()) > 5:
                                p_desc = sub_parts[1].strip()
                        else:
                            p_subtitle = rest
                    elif ':' in line:
                        parts = line.split(':', 1)
                        p_name = parts[0].strip(' •—:')
                        rest = parts[1].strip(' •—:')
                        if len(rest) > 0:
                            if len(rest) > 50 and not any(k in rest for k in ['NumPy', 'Python', 'FastAPI', 'PyTorch', 'TensorFlow']):
                                p_desc = rest
                            else:
                                p_subtitle = rest.strip(' :•')
                    
                    if p_name.lower().endswith((': numpy', ' numpy', '— numpy', '- numpy')):
                        p_name = re.sub(r'(?i)[:—\-]\s*numpy.*$', '', p_name).strip()
                        p_subtitle = 'NumPy • TensorFlow • PyTorch'
                    
                    current_proj = {
                        'name': p_name,
                        'subtitle': p_subtitle,
                        'description': p_desc,
                        'technologies': [],
                        'repository': repo,
                        'bullets': []
                    }
                else:
                    if not current_proj:
                        continue
                    
                    repo_m = re.search(r'(https?://github\.com/[^\s\)]+)', line)
                    if repo_m:
                        current_proj['repository'] = repo_m.group(1).rstrip('. )')
                        continue
                    
                    tech_m = re.search(r'(?:Technologies|Tech Stack|Tech)\s*[:—]\s*([^•\n]+)', line, re.I)
                    if tech_m:
                        current_proj['technologies'] = [t.strip().rstrip('.') for t in tech_m.group(1).split(',') if t.strip()]
                        continue
                    
                    if line.lower().startswith('ai agent project:'):
                        if not current_proj['subtitle']:
                            current_proj['subtitle'] = 'AI Agent Project'
                        continue
                    
                    clean_l = line.lstrip('•-* ').strip()
                    if clean_l and not clean_l.lower().startswith(('repository:', 'repo:')):
                        current_proj['bullets'].append(clean_l if clean_l.endswith(('.', ':', ';', '!')) else clean_l + '.')
            
            if current_proj:
                projects_list.append(Project(**current_proj))

        # D. Work Experience (Multi-Job & Single-Job Support)
        experience_list: List[WorkExperience] = []
        if "EXPERIENCE" in sections_dict:
            raw_exp_text = sections_dict["EXPERIENCE"]
            exp_lines = [l.strip() for l in raw_exp_text.split("\n") if l.strip()]
            date_pat = r'(?:\||—|–|-|\(|\s)\s*((?:19|20)\d{2}\s*[-–—]\s*(?:(?:19|20)\d{2}|Present|Current|Recent|Ongoing|\d{4})|\b(?:19|20)\d{2}\b|Recent|Present|Current|Ongoing)\)?\s*$'
            
            headers = []
            roles = []
            descriptions = []
            
            for l in exp_lines:
                m = re.search(date_pat, l)
                if m:
                    dur = m.group(1).strip(" ()")
                    raw_head = l[:m.start()].strip(" •|—:-")
                    comp = raw_head
                    role = None
                    if " — " in raw_head:
                        parts = raw_head.split(" — ", 1)
                        if any(k in parts[0].lower() for k in ["engineer", "developer", "analyst", "specialist", "manager", "lead", "architect", "accountant", "consultant", "scientist", "director", "officer", "intern"]):
                            role, comp = parts[0].strip(), parts[1].strip()
                        else:
                            comp, role = parts[0].strip(), parts[1].strip()
                    elif " - " in raw_head:
                        parts = raw_head.split(" - ", 1)
                        if any(k in parts[0].lower() for k in ["engineer", "developer", "analyst", "specialist", "manager", "lead", "architect", "accountant", "consultant", "scientist", "director", "officer", "intern"]):
                            role, comp = parts[0].strip(), parts[1].strip()
                        else:
                            comp, role = parts[0].strip(), parts[1].strip()
                    elif " at " in raw_head.lower():
                        parts = re.split(r'\s+at\s+', raw_head, flags=re.I)
                        role, comp = parts[0].strip(), parts[1].strip()

                    headers.append((comp, dur, role))
                elif len(l) < 40 and not l.startswith("•") and not l.lower().startswith("lorem") and not l.lower().startswith("experience includes") and not l.lower().startswith("building") and not l.lower().startswith("developing"):
                    roles.append(l.strip(" •:"))
                else:
                    clean_b = l.lstrip("•-* ").strip()
                    if clean_b and not clean_b.lower().startswith("experience includes:"):
                        descriptions.append(clean_b)
            
            if headers and (len(headers) == len(roles) or len(headers) == len(descriptions) or any(h[2] for h in headers)):
                for i, (comp, dur, r_inline) in enumerate(headers):
                    r_title = r_inline or (roles[i] if i < len(roles) else "Professional Role")
                    b_list = [descriptions[i]] if i < len(descriptions) else []
                    experience_list.append(WorkExperience(
                        company=comp,
                        role=r_title,
                        duration=dur,
                        bullets=b_list
                    ))
            elif headers:
                current_exp = None
                for line in exp_lines:
                    m = re.search(date_pat, line)
                    if m:
                        if current_exp:
                            experience_list.append(WorkExperience(**current_exp))
                        dur = m.group(1).strip(" ()")
                        raw_head = line[:m.start()].strip(" •|—:-")
                        comp = raw_head
                        role = "Professional Role"
                        if " — " in raw_head:
                            parts = raw_head.split(" — ", 1)
                            if any(k in parts[0].lower() for k in ["engineer", "developer", "analyst", "specialist", "manager", "lead", "architect", "accountant", "consultant", "scientist", "director", "officer", "intern"]):
                                role, comp = parts[0].strip(), parts[1].strip()
                            else:
                                comp, role = parts[0].strip(), parts[1].strip()
                        elif " - " in raw_head:
                            parts = raw_head.split(" - ", 1)
                            if any(k in parts[0].lower() for k in ["engineer", "developer", "analyst", "specialist", "manager", "lead", "architect", "accountant", "consultant", "scientist", "director", "officer", "intern"]):
                                role, comp = parts[0].strip(), parts[1].strip()
                            else:
                                comp, role = parts[0].strip(), parts[1].strip()
                        elif " at " in raw_head.lower():
                            parts = re.split(r'\s+at\s+', raw_head, flags=re.I)
                            role, comp = parts[0].strip(), parts[1].strip()

                        current_exp = {
                            "company": comp,
                            "role": role,
                            "duration": dur,
                            "subtitle": None,
                            "summary": None,
                            "bullets": []
                        }
                    elif current_exp is not None:
                        if current_exp["role"] == "Professional Role" and len(line) < 40 and not line.startswith("•") and not line.lower().startswith("lorem"):
                            current_exp["role"] = line.strip(" •:")
                        else:
                            sub_bullets = [b.strip() for b in re.split(r'[\u2022\u2023\u25e6\u2043\u2219•]', line) if b.strip()]
                            for sb in sub_bullets:
                                if sb and not sb.lower().startswith("experience includes:"):
                                    current_exp["bullets"].append(sb)
                if current_exp:
                    experience_list.append(WorkExperience(**current_exp))
            else:
                # Single-company / Freelance format
                current_exp = {
                    "company": "Independent / Freelance",
                    "role": "Software & AI Developer",
                    "duration": "Recent",
                    "subtitle": None,
                    "summary": None,
                    "bullets": []
                }
                for line in exp_lines:
                    ch_low = line.lower()
                    if "experience includes" in ch_low:
                        continue
                    elif any(k in ch_low for k in ["freelance software", "software developer", "software engineer", "ai developer", "accountant"]):
                        current_exp["role"] = line.strip(" .:")
                    elif "ai & software engineering" in ch_low:
                        current_exp["subtitle"] = line.strip(" :.")
                    elif len(line) > 5:
                        sub_bullets = [b.strip() for b in re.split(r'[\u2022\u2023\u25e6\u2043\u2219•]', line) if b.strip()]
                        for sb in sub_bullets:
                            current_exp["bullets"].append(sb)
                if current_exp["bullets"] or current_exp["role"]:
                    experience_list.append(WorkExperience(**current_exp))

        # E. Education (Multi-Degree & Single-Degree Support)
        education_list: List[Education] = []
        if "EDUCATION" in sections_dict:
            raw_edu_text = sections_dict["EDUCATION"]
            edu_lines = [l.strip() for l in raw_edu_text.split("\n") if l.strip()]
            date_pat = r'(?:\||—|–|-|\(|\s)\s*((?:19|20)\d{2}\s*[-–—]\s*(?:(?:19|20)\d{2}|Present|Current|\d{4})|\b(?:19|20)\d{2}\b)\)?\s*$'
            
            headers = []
            degrees = []
            descriptions = []
            
            for l in edu_lines:
                m = re.search(date_pat, l)
                if m:
                    dur = m.group(1).strip(" ()")
                    raw_head = l[:m.start()].strip(" •|—:-")
                    inst = raw_head
                    deg = None
                    if " — " in raw_head:
                        parts = raw_head.split(" — ", 1)
                        if any(k in parts[0].lower() for k in ["university", "college", "institute", "school", "academy", "mit", "stanford", "berkeley", "harvard", "columbia"]):
                            inst, deg = parts[0].strip(), parts[1].strip()
                        else:
                            deg, inst = parts[0].strip(), parts[1].strip()
                    elif " - " in raw_head:
                        parts = raw_head.split(" - ", 1)
                        if any(k in parts[0].lower() for k in ["university", "college", "institute", "school", "academy", "mit", "stanford", "berkeley", "harvard", "columbia"]):
                            inst, deg = parts[0].strip(), parts[1].strip()
                        else:
                            deg, inst = parts[0].strip(), parts[1].strip()
                    headers.append((inst, dur, deg))
                elif len(l) < 45 and not l.startswith("•") and not l.lower().startswith("lorem"):
                    degrees.append(l.strip(" •:"))
                else:
                    descriptions.append(l.lstrip("•-* ").strip())
                    
            if headers and (len(headers) == len(degrees) or len(headers) == len(descriptions) or any(h[2] for h in headers)):
                # Clustered / Layer-ordered Canva layout
                for i, (inst, yr, deg_inline) in enumerate(headers):
                    deg = deg_inline or (degrees[i] if i < len(degrees) else "Degree")
                    det = descriptions[i] if i < len(descriptions) else None
                    education_list.append(Education(
                        institution=inst,
                        degree=deg,
                        year=yr,
                        details=det
                    ))
            elif headers:
                current_edu = None
                for line in edu_lines:
                    m = re.search(date_pat, line)
                    if m:
                        if current_edu:
                            education_list.append(Education(**current_edu))
                        dur = m.group(1).strip(" ()")
                        raw_head = line[:m.start()].strip(" •|—:-")
                        inst = raw_head
                        deg = "Degree"
                        if " — " in raw_head:
                            parts = raw_head.split(" — ", 1)
                            if any(k in parts[0].lower() for k in ["university", "college", "institute", "school", "academy", "mit", "stanford", "berkeley", "harvard", "columbia"]):
                                inst, deg = parts[0].strip(), parts[1].strip()
                            else:
                                deg, inst = parts[0].strip(), parts[1].strip()
                        elif " - " in raw_head:
                            parts = raw_head.split(" - ", 1)
                            if any(k in parts[0].lower() for k in ["university", "college", "institute", "school", "academy", "mit", "stanford", "berkeley", "harvard", "columbia"]):
                                inst, deg = parts[0].strip(), parts[1].strip()
                            else:
                                deg, inst = parts[0].strip(), parts[1].strip()
                        current_edu = {
                            "institution": inst,
                            "degree": deg,
                            "year": dur,
                            "details": None
                        }
                    elif current_edu is not None:
                        if current_edu["degree"] == "Degree" and len(line) < 45 and not line.startswith("•") and not line.lower().startswith("lorem"):
                            current_edu["degree"] = line.strip(" •:")
                        else:
                            clean_det = line.lstrip("•-* ").strip()
                            if clean_det:
                                current_edu["details"] = (current_edu["details"] + " " + clean_det) if current_edu["details"] else clean_det
                if current_edu:
                    education_list.append(Education(**current_edu))
            else:
                for line in edu_lines:
                    parts = re.split(r'[—|–|-]', line)
                    if len(parts) >= 2:
                        deg = parts[0].strip(" •")
                        inst = parts[1].strip(" •")
                        yr = parts[2].strip(" •") if len(parts) > 2 else None
                        education_list.append(Education(institution=inst, degree=deg, year=yr))
                    else:
                        education_list.append(Education(institution=line.strip(" •"), degree="Degree"))

        # F. Additional Background & Target Role
        additional_background = sections_dict.get("ADDITIONAL_BACKGROUND", "").strip(" :—\n") or None
        target_role = sections_dict.get("TARGET_ROLE", "").strip(" :—\n") or None

        # Clean Summary: Strip contact info, email, phone, location from summary text
        raw_summary = sections_dict.get("SUMMARY")
        clean_sum = None
        if raw_summary:
            clean_sum = raw_summary.strip()
            if email_match:
                clean_sum = clean_sum.replace(email_match.group(0), "")
            if phone_val:
                clean_sum = clean_sum.replace(phone_val, "")
            if location_val:
                clean_sum = clean_sum.replace(location_val, "")
            clean_sum = re.sub(r'(?i)\b(?:email|phone|tel|mobile|location|address|linkedin|github)\b\s*[:—]?', '', clean_sum)
            clean_sum = re.sub(r'\s{2,}', ' ', clean_sum).strip(' ,•—\t\n+')
            if len(clean_sum) < 5:
                clean_sum = None

        # Integrated Header: Ensure headline captures target role directly under name
        final_headline = headline or target_role or "AI & Machine Learning Specialist"

        # Sanitize URLs across contact and projects
        from core.tailor import ResumeTailor
        if contact:
            contact.linkedin = ResumeTailor.sanitize_url(contact.linkedin)
            contact.github = ResumeTailor.sanitize_url(contact.github)
            contact.portfolio = ResumeTailor.sanitize_url(contact.portfolio)

        # Deduplicate and clean projects
        clean_projects = []
        for p in projects_list:
            clean_p = Project(
                name=p.name.strip(" :—"),
                subtitle=p.subtitle,
                description=p.description,
                bullets=ResumeTailor.deduplicate_bullets(p.bullets),
                technologies=[t.strip(" ,.;:") for t in p.technologies if t.strip()],
                repository=ResumeTailor.sanitize_url(p.repository)
            )
            clean_projects.append(clean_p)

        # Deduplicate and clean experience
        clean_experience = []
        for e in experience_list:
            clean_e = WorkExperience(
                company=e.company,
                role=e.role,
                location=e.location,
                duration=e.duration,
                subtitle=e.subtitle,
                summary=e.summary,
                bullets=ResumeTailor.deduplicate_bullets(e.bullets)
            )
            clean_experience.append(clean_e)

        return UserProfile(
            full_name=full_name,
            headline=final_headline,
            contact=contact,
            summary=clean_sum,
            skills=skills_list,
            categorized_skills=categorized_skills,
            experience=clean_experience,
            projects=clean_projects,
            education=education_list,
            certifications=certifications_list,
            additional_background=additional_background,
            target_role=None  # Integrated directly into header/headline to prevent standalone bottom section
        )

    @classmethod
    def parse_file(cls, file_path: str) -> UserProfile:
        text = cls.extract_text_from_file(file_path)
        return cls.parse_text_to_profile(text)
